from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from flask_admin import Admin, BaseView, expose
from flask_admin.contrib.sqla import ModelView
from flask_admin.form import ImageUploadField
from wtforms.fields import SelectField
import os
from io import BytesIO
import openpyxl
from openpyxl.utils import get_column_letter
from reportlab.pdfgen import canvas
from docx import Document

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///archery.db'
app.config['SECRET_KEY'] = 'your-secret-key'
db = SQLAlchemy(app)

UPLOAD_PATH = os.path.join(os.path.dirname(__file__), 'static/uploads')
os.makedirs(UPLOAD_PATH, exist_ok=True)

class News(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    content = db.Column(db.Text, nullable=False)
    date = db.Column(db.DateTime, default=datetime.utcnow)
    category = db.Column(db.String(50))
    image_url = db.Column(db.String(200))
    image = db.Column(db.String(100)) 

class Event(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, nullable=False)
    date = db.Column(db.DateTime, nullable=False)
    location = db.Column(db.String(200))
    contact = db.Column(db.String(200))
    image_url = db.Column(db.String(200))
    image = db.Column(db.String(100))
    gender_category = db.Column(db.String(10))  # Мужской / Женский / Оба

    participants = db.relationship('Participant', back_populates='event', lazy=True, cascade='all, delete-orphan')

    def __str__(self):
        return self.title

class Participant(db.Model):
    
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    phone = db.Column(db.String(20))
    email = db.Column(db.String(100))
    event_id = db.Column(db.Integer, db.ForeignKey('event.id'), nullable=False)
    event = db.relationship('Event', back_populates='participants')

class NewsModelView(ModelView):
    form_extra_fields = {
        'image': ImageUploadField(
            'Изображение',
            base_path=UPLOAD_PATH,
            relative_path='',
            url_relative_path='/static/uploads/'
        )
    }

class EventModelView(ModelView):
    form_extra_fields = {
        'image': ImageUploadField(
            'Изображение',
            base_path=UPLOAD_PATH,
            relative_path='',
            url_relative_path='/static/uploads/'
        ),
        'gender_category': SelectField(
            'Категория по полу',
            choices=[('Мужской', 'Мужской'), ('Женский', 'Женский'), ('Оба', 'Оба')]
        )
    }

class ParticipantModelView(ModelView):
    column_list = ('name', 'email', 'phone', 'event')
    column_labels = {'event': 'Соревнование'}
    column_filters = ['event']
    column_searchable_list = ['name', 'email']
    column_sortable_list = ['name', 'email', 'event']
    can_create = False

# 📄 Экспорт в Excel
class ExportParticipantsExcel(BaseView):
    @expose('/')
    def index(self):
        participants = Participant.query.all()

        # Создание Excel-файла
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Participants"
        
        # Заголовки
        headers = ['Имя', 'Email', 'Телефон', 'Соревнование']
        for col_num, header in enumerate(headers, 1):
            ws[f"{get_column_letter(col_num)}1"] = header

        # Заполнение данными
        for row_num, participant in enumerate(participants, 2):
            ws[f"A{row_num}"] = participant.name
            ws[f"B{row_num}"] = participant.email
            ws[f"C{row_num}"] = participant.phone
            ws[f"D{row_num}"] = participant.event.title

        # Сохранение в буфер
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        return send_file(buffer, as_attachment=True, download_name="participants.xlsx", mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


# 📄 Экспорт в DOCX
class ExportParticipantsDOCX(BaseView):
    @expose('/')
    def index(self):
        participants = Participant.query.all()

        # Создание документа Word
        doc = Document()
        doc.add_heading('Список участников', 0)

        # Заголовки
        doc.add_paragraph('Имя | Email | Телефон | Соревнование')

        # Заполнение данными
        for participant in participants:
            doc.add_paragraph(f"{participant.name}    | {    participant.email}    |   {participant.phone}    |     {participant.event.title}")

        # Сохранение в буфер
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(buffer, as_attachment=True, download_name="participants.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# 🛠 Admin Panel
admin = Admin(app, name='Archery Federation Admin', template_mode='bootstrap3')
admin.add_view(NewsModelView(News, db.session))
admin.add_view(EventModelView(Event, db.session))
admin.add_view(ParticipantModelView(Participant, db.session))
admin.add_view(ExportParticipantsExcel(name='Экспорт в Excel', endpoint='export_excel'))
admin.add_view(ExportParticipantsDOCX(name='Экспорт в DOCX', endpoint='export_docx'))


@app.route('/')
def home():
    latest_news = News.query.order_by(News.date.desc()).limit(3).all()
    upcoming_events = Event.query.filter(Event.date >= datetime.utcnow()).order_by(Event.date).limit(3).all()
    return render_template('index.html', news=latest_news, events=upcoming_events)


@app.route('/news')
def news():
    category = request.args.get('category')
    if category:
        news_items = News.query.filter_by(category=category).order_by(News.date.desc()).all()
    else:
        news_items = News.query.order_by(News.date.desc()).all()
    return render_template('news.html', news=news_items)


@app.route('/news/<int:id>')
def news_detail(id):
    news_item = News.query.get_or_404(id)
    return render_template('news_detail.html', news=news_item)


@app.route('/events')
def events():
    gender = request.args.get('gender')
    query = Event.query
    if gender in ['Мужской', 'Женский']:
        query = query.filter_by(gender_category=gender)
    events = query.order_by(Event.date.desc()).all()
    return render_template('events.html', events=events)


@app.route('/register/<int:event_id>', methods=['GET', 'POST'])
def register(event_id):
    event = Event.query.get_or_404(event_id)
    if request.method == 'POST':
        name = request.form['name']
        phone = request.form['phone']
        email = request.form['email']
        participant = Participant(name=name, phone=phone, email=email, event=event)
        db.session.add(participant)
        db.session.commit()
        flash('Вы успешно зарегистрировались на мероприятие!', 'success')
        return redirect(url_for('events'))
    return render_template('register.html', event=event)


@app.route('/contact')
def contact():
    return render_template('contact.html')


if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)
